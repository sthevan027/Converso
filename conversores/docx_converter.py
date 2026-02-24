from __future__ import annotations

import io
import re
import tempfile
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as create_document
from docx.document import Document as WordDocument
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Cm

from utils.ocr_engine import is_scanned_page, ocr_page, OCRLine

from .base import (
    BaseConverter,
    ConversionOptions,
    ConversionResult,
    HeaderFooterMode,
)


def _detect_document_type(pdf_path: str) -> str:
    """Detecta se o PDF é nativo, escaneado ou misto."""
    doc = fitz.open(pdf_path)
    try:
        native_count = 0
        scanned_count = 0
        for page in doc:
            if is_scanned_page(page):
                scanned_count += 1
            else:
                native_count += 1

        if scanned_count == 0:
            return "native"
        if native_count == 0:
            return "scanned"
        return "mixed"
    finally:
        doc.close()


class DocxConverter(BaseConverter):
    """Conversor robusto de PDF para DOCX.

    - PDFs nativos: usa pdf2docx para máxima fidelidade.
    - PDFs escaneados: usa OCR (RapidOCR) + construção estruturada.
    - PDFs mistos: combina ambas as abordagens.
    """

    def __init__(self, options: Optional[ConversionOptions] = None) -> None:
        super().__init__(options)
        self._result: ConversionResult = None  # type: ignore

    def convert(self, pdf_path: str, output_path: str) -> ConversionResult:
        pdf_file = Path(pdf_path)
        if not pdf_file.is_file():
            raise FileNotFoundError(f"Arquivo PDF não encontrado: {pdf_file}")

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self._result = ConversionResult(success=False, output_path=str(output_file))

        start, end = self._resolve_page_range()
        doc_type = _detect_document_type(str(pdf_file))

        self._log(f"[DOCX] Tipo detectado: {doc_type} | '{pdf_file.name}'")

        try:
            if doc_type == "native":
                self._convert_native(str(pdf_file), str(output_file), start, end)
            elif doc_type == "scanned":
                self._convert_scanned(str(pdf_file), str(output_file), start, end)
            else:
                self._convert_mixed(str(pdf_file), str(output_file), start, end)

            self._result.success = True
            self._log(f"[DOCX] Arquivo gerado: {output_file}")
        except Exception as e:
            self._result.errors.append(str(e))
            raise

        return self._result

    # ------------------------------------------------------------------
    # PDFs nativos → pdf2docx (alta fidelidade)
    # ------------------------------------------------------------------

    def _convert_native(
        self, pdf_path: str, output_path: str,
        start: Optional[int], end: Optional[int],
    ) -> None:
        from pdf2docx import Converter as Pdf2DocxConverter

        self._log("[DOCX] Convertendo PDF nativo via pdf2docx...")

        cv = Pdf2DocxConverter(pdf_path)
        try:
            pages = self._build_page_list(pdf_path, start, end)
            cv.convert(output_path, pages=pages)
            self._result.pages_converted = len(pages) if pages else self._count_pages(pdf_path)
        finally:
            cv.close()

    # ------------------------------------------------------------------
    # PDFs escaneados → OCR + DOCX estruturado
    # ------------------------------------------------------------------

    def _convert_scanned(
        self, pdf_path: str, output_path: str,
        start: Optional[int], end: Optional[int],
    ) -> None:
        self._log("[DOCX] Convertendo PDF escaneado via OCR...")

        doc = fitz.open(pdf_path)
        try:
            word_doc = create_document()
            pages = self._get_page_range(doc, start, end)

            first = True
            for page_num in pages:
                if not first:
                    word_doc.add_page_break()
                first = False

                page = doc[page_num]
                self._setup_section_from_page(word_doc, page, page_num == pages[0])
                self._ocr_page_to_docx(word_doc, page, page_num)
                self._result.pages_converted += 1

            word_doc.save(output_path)
        finally:
            doc.close()

    # ------------------------------------------------------------------
    # PDFs mistos → pdf2docx para nativas + OCR para escaneadas
    # ------------------------------------------------------------------

    def _convert_mixed(
        self, pdf_path: str, output_path: str,
        start: Optional[int], end: Optional[int],
    ) -> None:
        self._log("[DOCX] Convertendo PDF misto (nativo + escaneado)...")

        doc = fitz.open(pdf_path)
        try:
            pages = self._get_page_range(doc, start, end)

            native_pages = []
            scanned_pages = []
            for pn in pages:
                if is_scanned_page(doc[pn]):
                    scanned_pages.append(pn)
                else:
                    native_pages.append(pn)

            self._log(f"  Nativas: {len(native_pages)} | Escaneadas: {len(scanned_pages)}")

            if native_pages and not scanned_pages:
                doc.close()
                self._convert_native(pdf_path, output_path, start, end)
                return
            if scanned_pages and not native_pages:
                doc.close()
                self._convert_scanned(pdf_path, output_path, start, end)
                return

            native_tmp = None
            try:
                if native_pages:
                    from pdf2docx import Converter as Pdf2DocxConverter
                    native_tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
                    native_tmp.close()
                    cv = Pdf2DocxConverter(pdf_path)
                    cv.convert(native_tmp.name, pages=native_pages)
                    cv.close()

                word_doc = create_document()

                first = True
                for page_num in pages:
                    if not first:
                        word_doc.add_page_break()
                    first = False

                    page = doc[page_num]
                    self._setup_section_from_page(word_doc, page, page_num == pages[0])

                    if page_num in scanned_pages:
                        self._ocr_page_to_docx(word_doc, page, page_num)
                    else:
                        self._native_page_to_docx(word_doc, page, page_num)

                    self._result.pages_converted += 1

                word_doc.save(output_path)
            finally:
                if native_tmp:
                    Path(native_tmp.name).unlink(missing_ok=True)
        finally:
            doc.close()

    # ------------------------------------------------------------------
    # OCR de uma página → parágrafos no DOCX
    # ------------------------------------------------------------------

    def _ocr_page_to_docx(
        self, word_doc: WordDocument, page: fitz.Page, page_num: int,
    ) -> None:
        self._log(f"  [OCR] Página {page_num + 1}...")
        ocr_result = ocr_page(page, dpi=300)

        if not ocr_result.text_lines:
            self._result.warnings.append(f"Página {page_num + 1}: nenhum texto detectado pelo OCR")
            return

        groups = self._group_ocr_lines(ocr_result.text_lines, page.rect.height)

        for group in groups:
            text = " ".join(line.text for line in group)
            text = self._normalize_ocr_spacing(text)

            if not text.strip():
                continue

            para = word_doc.add_paragraph()

            avg_height = sum(l.bbox[3] - l.bbox[1] for l in group) / len(group)
            font_size = max(min(avg_height * 0.65, 24.0), 8.0)

            if font_size > 13 and len(text) < 80:
                if font_size >= 18:
                    para.style = "Heading 1"
                elif font_size >= 14:
                    para.style = "Heading 2"
                else:
                    para.style = "Heading 3"
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = para.add_run(text)
            run.font.size = Pt(min(font_size, 12))

    def _group_ocr_lines(
        self, lines: list[OCRLine], page_height: float,
    ) -> list[list[OCRLine]]:
        """Agrupa linhas OCR em parágrafos baseado em proximidade vertical."""
        if not lines:
            return []

        sorted_lines = sorted(lines, key=lambda l: (l.bbox[1], l.bbox[0]))

        groups: list[list[OCRLine]] = []
        current: list[OCRLine] = [sorted_lines[0]]

        for line in sorted_lines[1:]:
            prev = current[-1]
            prev_bottom = prev.bbox[3]
            curr_top = line.bbox[1]
            gap = curr_top - prev_bottom

            prev_height = prev.bbox[3] - prev.bbox[1]
            curr_height = line.bbox[3] - line.bbox[1]
            avg_height = (prev_height + curr_height) / 2

            if gap > avg_height * 1.5:
                groups.append(current)
                current = [line]
            else:
                current.append(line)

        if current:
            groups.append(current)

        return groups

    @staticmethod
    def _normalize_ocr_spacing(text: str) -> str:
        """Corrige espaçamentos típicos de OCR sem quebrar palavras."""
        text = re.sub(r"(\w):(\w)", r"\1: \2", text)
        text = re.sub(r"(\w)\.(\s?\w)", r"\1. \2", text)
        text = re.sub(r"([a-záéíóúâêôãõüç])([A-ZÁÉÍÓÚÂÊÔÃÕÜÇ])", r"\1 \2", text)
        text = re.sub(r"\s{2,}", " ", text)
        return text.strip()

    # ------------------------------------------------------------------
    # Página nativa (texto) → parágrafos no DOCX
    # ------------------------------------------------------------------

    def _native_page_to_docx(
        self, word_doc: WordDocument, page: fitz.Page, page_num: int,
    ) -> None:
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

        for block in blocks:
            if block["type"] == 1 and self.options.extract_images:
                self._add_image_block(word_doc, page, block)
                continue

            if block["type"] != 0:
                continue

            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue

                line_text = "".join(s.get("text", "") for s in spans).strip()
                if not line_text:
                    continue

                para = word_doc.add_paragraph()

                for span in spans:
                    span_text = span.get("text", "")
                    if not span_text:
                        continue

                    run = para.add_run(span_text)
                    size = span.get("size", 11)
                    flags = span.get("flags", 0)

                    run.font.size = Pt(min(size, 24))
                    run.bold = bool(flags & 2**4)
                    run.italic = bool(flags & 2**1)

    def _add_image_block(
        self, word_doc: WordDocument, page: fitz.Page, block: dict,
    ) -> None:
        try:
            bbox = block.get("bbox", (0, 0, 0, 0))
            img_width_pt = bbox[2] - bbox[0]

            xref = block.get("image", None)
            if xref is None:
                img_list = page.get_images(full=True)
                if not img_list:
                    return
                xref = img_list[0][0]

            base_image = page.parent.extract_image(xref)
            if not base_image:
                return

            image_stream = io.BytesIO(base_image["image"])
            max_w = min(Inches(self.options.max_image_width / 96), Pt(img_width_pt))
            word_doc.add_picture(image_stream, width=max_w)
            self._result.images_extracted += 1
        except Exception as e:
            self._result.warnings.append(f"Erro ao extrair imagem: {e}")

    # ------------------------------------------------------------------
    # Configuração de seção / layout
    # ------------------------------------------------------------------

    def _setup_section_from_page(
        self, word_doc: WordDocument, page: fitz.Page, is_first: bool,
    ) -> None:
        w, h = page.rect.width, page.rect.height

        if is_first:
            section = word_doc.sections[0]
        else:
            section = word_doc.add_section()

        if w > h:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Pt(h)
            section.page_height = Pt(w)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Pt(w)
            section.page_height = Pt(h)

        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    # ------------------------------------------------------------------
    # Utilitários
    # ------------------------------------------------------------------

    def _resolve_page_range(self) -> tuple[Optional[int], Optional[int]]:
        start = self.options.start_page
        end = self.options.end_page

        if start is not None and start <= 0:
            raise ValueError("start_page deve ser >= 1")
        if end is not None and end <= 0:
            raise ValueError("end_page deve ser >= 1")
        if start is not None and end is not None and end < start:
            raise ValueError("end_page não pode ser menor que start_page")

        return start, end

    def _get_page_range(
        self, doc: fitz.Document,
        start: Optional[int], end: Optional[int],
    ) -> list[int]:
        first = (start - 1) if start else 0
        last = end if end else len(doc)
        return list(range(first, min(last, len(doc))))

    def _build_page_list(
        self, pdf_path: str,
        start: Optional[int], end: Optional[int],
    ) -> list[int] | None:
        if start is None and end is None:
            return None
        doc = fitz.open(pdf_path)
        try:
            return self._get_page_range(doc, start, end)
        finally:
            doc.close()

    @staticmethod
    def _count_pages(pdf_path: str) -> int:
        doc = fitz.open(pdf_path)
        try:
            return len(doc)
        finally:
            doc.close()
