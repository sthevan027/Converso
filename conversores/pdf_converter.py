from __future__ import annotations

from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt

from .base import BaseConverter, ConversionOptions, ConversionResult


class PdfConverter(BaseConverter):
    """Conversor de DOCX para PDF usando PyMuPDF."""

    def __init__(self, options: Optional[ConversionOptions] = None) -> None:
        super().__init__(options)
        self._result: ConversionResult = None  # type: ignore

    def convert(self, input_path: str, output_path: str) -> ConversionResult:
        input_file = Path(input_path)
        if not input_file.is_file():
            raise FileNotFoundError(f"Arquivo não encontrado: {input_file}")

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self._result = ConversionResult(
            success=False,
            output_path=str(output_file),
        )

        ext = input_file.suffix.lower()

        self._log(f"[PDF] Convertendo '{input_file}' -> '{output_file}'...")

        try:
            if ext == ".docx":
                self._convert_docx_to_pdf(str(input_file), str(output_file))
            elif ext == ".txt":
                self._convert_txt_to_pdf(str(input_file), str(output_file))
            elif ext == ".md":
                self._convert_md_to_pdf(str(input_file), str(output_file))
            else:
                raise ValueError(f"Formato de entrada não suportado: {ext}")

            self._result.success = True
            self._log(f"[PDF] Arquivo gerado: {output_file}")

        except Exception as e:
            self._result.errors.append(str(e))
            raise

        return self._result

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """Converte DOCX para PDF."""
        doc = Document(docx_path)
        pdf_doc = fitz.open()

        page_width = 595  # A4 width in points
        page_height = 842  # A4 height in points
        margin = 72  # 1 inch margin
        line_height = 14
        current_y = margin
        current_page = None

        def new_page():
            nonlocal current_page, current_y
            current_page = pdf_doc.new_page(width=page_width, height=page_height)
            current_y = margin
            return current_page

        current_page = new_page()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                current_y += line_height / 2
                continue

            font_size = 11
            font_name = "helv"
            is_bold = False

            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if "heading 1" in style_name:
                    font_size = 18
                    is_bold = True
                elif "heading 2" in style_name:
                    font_size = 14
                    is_bold = True
                elif "heading 3" in style_name:
                    font_size = 12
                    is_bold = True
                elif "title" in style_name:
                    font_size = 20
                    is_bold = True

            if para.runs:
                for run in para.runs:
                    if run.bold:
                        is_bold = True
                    if run.font.size:
                        font_size = run.font.size.pt

            line_height = font_size * 1.5
            text_width = page_width - 2 * margin

            lines = self._wrap_text(text, text_width, font_size)

            for line in lines:
                if current_y + line_height > page_height - margin:
                    current_page = new_page()
                    self._result.pages_converted += 1

                text_point = fitz.Point(margin, current_y)

                if is_bold:
                    current_page.insert_text(
                        text_point,
                        line,
                        fontsize=font_size,
                        fontname="helv",
                        color=(0, 0, 0),
                    )
                else:
                    current_page.insert_text(
                        text_point,
                        line,
                        fontsize=font_size,
                        fontname="helv",
                        color=(0, 0, 0),
                    )

                current_y += line_height

            current_y += line_height * 0.3

        for table in doc.tables:
            self._add_table_to_pdf(pdf_doc, table, current_page, margin, current_y)

        self._result.pages_converted = len(pdf_doc)
        pdf_doc.save(pdf_path)
        pdf_doc.close()

    def _wrap_text(self, text: str, max_width: float, font_size: float) -> list[str]:
        """Quebra texto em linhas que cabem na largura especificada."""
        chars_per_line = int(max_width / (font_size * 0.5))
        words = text.split()
        lines = []
        current_line = []
        current_length = 0

        for word in words:
            word_length = len(word)
            if current_length + word_length + 1 <= chars_per_line:
                current_line.append(word)
                current_length += word_length + 1
            else:
                if current_line:
                    lines.append(" ".join(current_line))
                current_line = [word]
                current_length = word_length

        if current_line:
            lines.append(" ".join(current_line))

        return lines if lines else [""]

    def _add_table_to_pdf(
        self,
        pdf_doc: fitz.Document,
        table,
        page: fitz.Page,
        margin: float,
        y_pos: float,
    ) -> None:
        """Adiciona tabela ao PDF (simplificado)."""
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                if y_pos > 750:
                    page = pdf_doc.new_page()
                    y_pos = margin

                page.insert_text(
                    fitz.Point(margin, y_pos),
                    row_text,
                    fontsize=10,
                )
                y_pos += 14

    def _convert_txt_to_pdf(self, txt_path: str, pdf_path: str) -> None:
        """Converte arquivo de texto para PDF."""
        with open(txt_path, "r", encoding="utf-8") as f:
            content = f.read()

        pdf_doc = fitz.open()
        page = pdf_doc.new_page(width=595, height=842)

        margin = 72
        current_y = margin
        font_size = 11
        line_height = font_size * 1.5

        lines = content.split("\n")

        for line in lines:
            if current_y + line_height > 770:
                page = pdf_doc.new_page(width=595, height=842)
                current_y = margin
                self._result.pages_converted += 1

            wrapped = self._wrap_text(line, 451, font_size)
            for wrapped_line in wrapped:
                page.insert_text(
                    fitz.Point(margin, current_y),
                    wrapped_line,
                    fontsize=font_size,
                )
                current_y += line_height

        self._result.pages_converted = len(pdf_doc)
        pdf_doc.save(pdf_path)
        pdf_doc.close()

    def _convert_md_to_pdf(self, md_path: str, pdf_path: str) -> None:
        """Converte Markdown para PDF (básico)."""
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()

        pdf_doc = fitz.open()
        page = pdf_doc.new_page(width=595, height=842)

        margin = 72
        current_y = margin

        lines = content.split("\n")

        for line in lines:
            font_size = 11
            is_heading = False

            if line.startswith("### "):
                font_size = 12
                line = line[4:]
                is_heading = True
            elif line.startswith("## "):
                font_size = 14
                line = line[3:]
                is_heading = True
            elif line.startswith("# "):
                font_size = 18
                line = line[2:]
                is_heading = True

            line_height = font_size * 1.5

            if current_y + line_height > 770:
                page = pdf_doc.new_page(width=595, height=842)
                current_y = margin
                self._result.pages_converted += 1

            if line.strip():
                page.insert_text(
                    fitz.Point(margin, current_y),
                    line,
                    fontsize=font_size,
                )

            current_y += line_height
            if is_heading:
                current_y += 5

        self._result.pages_converted = len(pdf_doc)
        pdf_doc.save(pdf_path)
        pdf_doc.close()
